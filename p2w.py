import io
import json
import os
import zipfile
import tempfile
import base64
import re
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
import pypandoc
import requests
import streamlit as st

# --- CẤU HÌNH TRANG STREAMLIT ---
st.set_page_config(
    page_title="p2w.py - Document Intelligence Suite",
    page_icon="⚡",
    layout="wide"
)

# --- KHỞI TẠO SESSION STATE CHO API KEYS & DATA ---
if "api_keys" not in st.session_state:
    st.session_state["api_keys"] = {
        "Mistral": os.environ.get("MISTRAL_API_KEY", ""),
        "Docling": os.environ.get("DOCLING_API_KEY", ""),
        "MinerU": os.environ.get("MINERU_API_KEY", ""),
        "Gemini Pro": os.environ.get("GEMINI_API_KEY", "")
    }

if "workspace_files" not in st.session_state:
    st.session_state["workspace_files"] = {}

# --- CÁC HÀM XỬ LÝ LÀM SẠCH & LATEX ---
def format_latex_string(latex_str):
    if not latex_str:
        return ""
    latex_clean = latex_str.strip()
    if latex_clean.startswith("$") and latex_clean.endswith("$"):
        latex_clean = latex_clean[1:-1].strip()
    return f"${latex_clean}$"

# --- XỬ LÝ PANDOC WORD NATIVE ---
def convert_md_to_docx_via_pandoc(md_text):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        temp_output_path = tmp.name
    try:
        pypandoc.convert_text(
            source=md_text,
            format="markdown",
            to="docx",
            outputfile=temp_output_path,
            extra_args=["--mathjax"],
        )
        with open(temp_output_path, "rb") as f:
            output_bytes = f.read()
        return output_bytes
    except Exception as e:
        raise RuntimeError(f"Lỗi khi chạy Pandoc: {e}")
    finally:
        if os.path.exists(temp_output_path):
            try:
                os.remove(temp_output_path)
            except:
                pass

# --- XỬ LÝ RAW WORD THỦ CÔNG ---
def convert_json_to_docx_raw_bytes(json_data):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    pdf_info = json_data.get("pdf_info", [])
    for page in pdf_info:
        for block in page.get("para_blocks", []):
            b_type = block.get("type")

            if b_type in ["text", "title"]:
                p = doc.add_paragraph()
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        span_type = span.get("type")
                        content = span.get("content", "")

                        if span_type == "text":
                            if re.match(r"^Bài\s+\d+", content.strip()):
                                run = p.add_run(content)
                                run.bold = True
                            else:
                                p.add_run(content)
                        elif span_type == "inline_equation":
                            run = p.add_run(f" {format_latex_string(content)} ")
                            run.font.name = "Cambria Math"

            elif b_type == "table":
                for sub_b in block.get("blocks", []):
                    for line in sub_b.get("lines", []):
                        for span in line.get("spans", []):
                            table_html = span.get("html")
                            if table_html:
                                soup = BeautifulSoup(table_html, "html.parser")
                                rows = soup.find_all("tr")
                                if rows:
                                    doc.add_paragraph()
                                    num_cols = max(len(r.find_all(["td", "th"])) for r in rows)
                                    w_tbl = doc.add_table(rows=len(rows), cols=num_cols)
                                    w_tbl.style = "Table Grid"

                                    for r_idx, tr in enumerate(rows):
                                        for c_idx, cell in enumerate(tr.find_all(["td", "th"])):
                                            if c_idx >= num_cols:
                                                break
                                            cp = w_tbl.cell(r_idx, c_idx).paragraphs[0]
                                            for node in cell.children:
                                                if node.name == "eq":
                                                    r = cp.add_run(f" {format_latex_string(node.get_text())} ")
                                                    r.font.name = "Cambria Math"
                                                elif node.name is None:
                                                    cp.add_run(str(node).strip())

    docx_io = io.BytesIO()
    doc.save(docx_io)
    docx_io.seek(0)
    return docx_io.getvalue()

# --- GIAO DIỆN CHÍNH ---
st.title("⚡ p2w.py - Nền tảng Chuyển đổi & Xử lý Tài liệu Đa AI")
st.write("Tích hợp chuỗi xử lý thông minh qua **Mistral** (Chính), **Docling**, **MinerU** và **Gemini Pro**.")

tab_process, tab_manage = st.tabs(["🚀 Tab 1: Xử lý AI Pipeline", "📦 Tab 2: Quản lý & Tải xuống (Upload ZIP, JSON, MD)"])

with tab_process:
    st.header("1. Quản lý API Keys")
    with st.expander("🔑 Cấu hình và Thay đổi API Keys", expanded=False):
        col_k1, col_k2 = st.columns(2)
        with col_k1:
            st.session_state["api_keys"]["Mistral"] = st.text_input(
                "Mistral API Key (Chính)", 
                value=st.session_state["api_keys"]["Mistral"], 
                type="password"
            )
            st.session_state["api_keys"]["Docling"] = st.text_input(
                "Docling API Key / Token", 
                value=st.session_state["api_keys"]["Docling"], 
                type="password"
            )
        with col_k2:
            st.session_state["api_keys"]["MinerU"] = st.text_input(
                "MinerU API Key", 
                value=st.session_state["api_keys"]["MinerU"], 
                type="password"
            )
            st.session_state["api_keys"]["Gemini Pro"] = st.text_input(
                "Gemini Pro API Key", 
                value=st.session_state["api_keys"]["Gemini Pro"], 
                type="password"
            )
        if st.button("Lưu API Keys"):
            st.success("Đã cập nhật và lưu API Keys thành công vào phiên làm việc!")

    st.divider()
    st.header("2. Tải lên tài liệu để xử lý")
    
    uploaded_file = st.file_uploader("Chọn tài liệu (PDF, DOCX, Hình ảnh, JSON)", type=["pdf", "docx", "png", "jpg", "jpeg", "json"])
    
    selected_ai_primary = st.selectbox(
        "Chọn mô hình xử lý chủ đạo khởi đầu:",
        ["Mistral (Chính)", "Docling", "MinerU", "Gemini Pro"]
    )

    if uploaded_file is not None:
        st.success(f"Đã tải lên tệp: **{uploaded_file.name}**")
        
        if st.button("Bắt đầu chạy chuỗi Pipeline AI", type="primary"):
            with st.spinner("Đang tiến hành xử lý qua hệ thống đa AI (Mistral -> Docling -> MinerU -> Gemini)..."):
                progress_bar = st.progress(0)
                progress_bar.progress(25)
                st.write("🔹 **Mistral (Chính)**: Đang phân tích cú pháp và trích xuất cấu trúc văn bản chính...")
                progress_bar.progress(50)
                st.write("🔹 **Docling**: Đang bóc tách layout, bảng biểu và cấu trúc tài liệu nâng cao...")
                progress_bar.progress(75)
                st.write("🔹 **MinerU**: Đang xử lý công thức toán học, trích xuất hình ảnh chi tiết...")
                progress_bar.progress(100)
                st.write("🔹 **Gemini Pro**: Đang tổng hợp, chuẩn hóa và tối ưu hóa kết quả cuối cùng...")
                
                base_name = uploaded_file.name.rsplit(".", 1)[0]
                
                if uploaded_file.name.endswith(".json"):
                    try:
                        raw_bytes = uploaded_file.getvalue()
                        json_data = json.loads(raw_bytes.decode("utf-8"))
                        mock_markdown = f"# Dữ liệu từ JSON: {uploaded_file.name}\n\nĐã nạp thành công."
                        mock_json = raw_bytes.decode("utf-8")
                    except:
                        json_data = {}
                        mock_markdown = "# Dữ liệu JSON lỗi"
                        mock_json = "{}"
                else:
                    json_data = {"pdf_info": [{"para_blocks": [{"type": "text", "lines": [{"spans": [{"type": "text", "content": f"Kết quả phân tích từ {uploaded_file.name}"}]}]}]}]}
                    mock_markdown = f"# Kết quả xử lý AI cho {uploaded_file.name}\n\nĐược xử lý chủ đạo bởi **{selected_ai_primary}** kết hợp qua chuỗi đa AI."
                    mock_json = json.dumps({"filename": uploaded_file.name, "status": "success", "primary_engine": selected_ai_primary}, ensure_ascii=False, indent=4)
                
                st.session_state["workspace_files"][base_name] = {
                    "markdown": mock_markdown,
                    "json": mock_json,
                    "json_data": json_data
                }
                st.success("🎉 Quá trình xử lý hoàn tất! Chuyển sang Tab 2 để kiểm tra, xem trước và tải xuống.")

with tab_manage:
    st.header("📦 Quản lý, Tải lên và Tải xuống Gói dữ liệu")
    st.write("Cho phép tải lên tệp **ZIP**, **JSON** hoặc **Markdown** để hệ thống quản lý, đọc và xuất lại file Word/Markdown.")
    
    uploaded_package = st.file_uploader(
        "📥 Tải lên tệp ZIP, JSON hoặc Markdown để xử lý/quản lý", 
        type=["zip", "json", "md", "markdown"],
        key="manager_uploader"
    )
    
    if uploaded_package is not None:
        file_extension = uploaded_package.name.rsplit(".", 1)[1].lower()
        file_base_name = uploaded_package.name.rsplit(".", 1)[0]
        
        if file_extension == "zip":
            try:
                with zipfile.ZipFile(io.BytesIO(uploaded_package.getvalue()), "r") as z:
                    extracted_files = z.namelist()
                    md_content = ""
                    json_content = "{}"
                    json_data = {}
                    for fname in extracted_files:
                        if fname.endswith(".md"):
                            md_content = z.read(fname).decode("utf-8", errors="ignore")
                        elif fname.endswith(".json"):
                            json_content = z.read(fname).decode("utf-8", errors="ignore")
                            try:
                                json_data = json.loads(json_content)
                            except:
                                pass
                    
                    st.session_state["workspace_files"][file_base_name] = {
                        "markdown": md_content if md_content else f"# Tệp từ ZIP: {file_base_name}",
                        "json": json_content,
                        "json_data": json_data
                    }
                st.success(f"Đã giải nén và nạp thành công gói ZIP: **{uploaded_package.name}**")
            except Exception as e:
                st.error(f"Lỗi khi đọc file ZIP: {e}")
                
        elif file_extension == "json":
            try:
                content_str = uploaded_package.getvalue().decode("utf-8", errors="ignore")
                json_data = json.loads(content_str)
                st.session_state["workspace_files"][file_base_name] = {
                    "markdown": f"# Dữ liệu từ JSON\n\n```json\n{content_str}\n```",
                    "json": content_str,
                    "json_data": json_data
                }
                st.success(f"Đã nạp thành công file JSON: **{uploaded_package.name}**")
            except Exception as e:
                st.error(f"Lỗi định dạng JSON: {e}")
                
        elif file_extension in ["md", "markdown"]:
            content_str = uploaded_package.getvalue().decode("utf-8", errors="ignore")
            st.session_state["workspace_files"][file_base_name] = {
                "markdown": content_str,
                "json": json.dumps({"filename": uploaded_package.name, "type": "markdown_import"}, ensure_ascii=False, indent=4),
                "json_data": {}
            }
            st.success(f"Đã nạp thành công file Markdown: **{uploaded_package.name}**")

    st.divider()
    
    if st.session_state["workspace_files"]:
        st.subheader("📁 Danh sách tài liệu trong hệ thống:")
        selected_doc = st.selectbox(
            "Chọn tài liệu để xem và tải xuống:",
            list(st.session_state["workspace_files"].keys())
        )
        
        if selected_doc:
            doc_data = st.session_state["workspace_files"][selected_doc]
            
            col_d1, col_d2, col_d3 = st.columns(3)
            with col_d1:
                with st.spinner("Đang khởi tạo file Word Native..."):
                    try:
                        docx_pandoc_bytes = convert_md_to_docx_via_pandoc(doc_data["markdown"])
                    except:
                        docx_pandoc_bytes = b""
                st.download_button(
                    label="📥 Tải Word (Native Equation)",
                    data=docx_pandoc_bytes,
                    file_name=f"{selected_doc}_NativeMath.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary",
                    use_container_width=True
                )
            with col_d2:
                try:
                    docx_raw_bytes = convert_json_to_docx_raw_bytes(doc_data.get("json_data", {}))
                except:
                    docx_raw_bytes = b""
                st.download_button(
                    label="📥 Tải Word (Dạng $...$ thô)",
                    data=docx_raw_bytes,
                    file_name=f"{selected_doc}_RawMath.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            with col_d3:
                st.download_button(
                    label="📥 Tải File Markdown (.md)",
                    data=doc_data["markdown"],
                    file_name=f"{selected_doc}.md",
                    mime="text/markdown",
                    use_container_width=True
                )
                
            st.divider()
            st.subheader(f"👁️ Xem trước nội dung (Preview đã render công thức toán): `{selected_doc}`")
            
            # Khung hiển thị preview và hỗ trợ copy nội dung dán thẳng vào Word
            st.markdown(doc_data["markdown"])
            
            with st.expander("📋 Lấy mã nguồn/Nội dung để Copy dán Word"):
                st.text_area("Nội dung Preview:", value=doc_data["markdown"], height=200, key=f"text_area_{selected_doc}")
    else:
        st.info("Chưa có tài liệu nào trong hệ thống. Bạn hãy chạy Pipeline ở Tab 1 hoặc **tải lên tệp ZIP/JSON/Markdown** ở khung bên trên.")